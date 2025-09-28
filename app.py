import streamlit as st
from PIL import Image
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import io
from datetime import datetime
import os
from dotenv import load_dotenv
import hashlib
import json
from supabase import create_client, Client

# Load environment variables
load_dotenv()

# Supabase connection
@st.cache_resource
def get_supabase_client():
    """Get Supabase client"""
    try:
        url = os.getenv('SUPABASE_URL')
        key = os.getenv('SUPABASE_ANON_KEY')
        
        if not url or not key:
            st.error("Supabase URL and key are required in environment variables")
            return None
        
        supabase: Client = create_client(url, key)
        return supabase
    except Exception as e:
        st.error(f"Supabase connection failed: {str(e)}")
        return None

def init_db():
    """Test Supabase connection and tables"""
    supabase = get_supabase_client()
    if not supabase:
        return False
    
    try:
        # Test if tables exist by trying to query them
        supabase.table('users').select('id').limit(1).execute()
        supabase.table('documents').select('id').limit(1).execute()
        return True
    except Exception as e:
        st.error(f"Database tables don't exist or connection failed: {str(e)}")
        st.info("Please create the required tables in your Supabase dashboard using the SQL provided in the setup instructions.")
        return False

def hash_password(password):
    """Hash password using SHA-256"""
    return hashlib.sha256(password.encode()).hexdigest()

def create_user(username, email, password):
    """Create a new user"""
    supabase = get_supabase_client()
    if not supabase:
        return None
    
    try:
        user_data = {
            "username": username,
            "email": email,
            "password_hash": hash_password(password),
            "created_at": datetime.now().isoformat()
        }
        
        result = supabase.table('users').insert(user_data).execute()
        
        if result.data:
            return result.data[0]['id']
        return None
    except Exception as e:
        if "duplicate key" in str(e).lower() or "unique" in str(e).lower():
            return None  # User already exists
        st.error(f"Error creating user: {str(e)}")
        return None

def authenticate_user(username, password):
    """Authenticate user login"""
    supabase = get_supabase_client()
    if not supabase:
        return None
    
    try:
        result = supabase.table('users').select('id, username').eq(
            'username', username
        ).eq(
            'password_hash', hash_password(password)
        ).execute()
        
        if result.data:
            user = result.data[0]
            return (user['id'], user['username'])
        return None
    except Exception as e:
        st.error(f"Authentication error: {str(e)}")
        return None

def save_document(user_id, doc_name, extracted_texts, image_names, settings):
    """Save document to database"""
    supabase = get_supabase_client()
    if not supabase:
        return None
    
    try:
        doc_data = {
            "user_id": user_id,
            "doc_name": doc_name,
            "extracted_texts": json.dumps(extracted_texts),
            "image_names": json.dumps(image_names),
            "processing_settings": json.dumps(settings),
            "created_at": datetime.now().isoformat()
        }
        
        result = supabase.table('documents').insert(doc_data).execute()
        
        if result.data:
            return result.data[0]['id']
        return None
    except Exception as e:
        st.error(f"Error saving document: {str(e)}")
        return None

def get_user_documents(user_id):
    """Get all documents for a user"""
    supabase = get_supabase_client()
    if not supabase:
        return []
    
    try:
        result = supabase.table('documents').select(
            'id, doc_name, extracted_texts, image_names, processing_settings, created_at'
        ).eq('user_id', user_id).order('created_at', desc=True).execute()
        
        docs = []
        for doc in result.data:
            docs.append((
                doc['id'],
                doc['doc_name'],
                doc['extracted_texts'],
                doc['image_names'],
                doc['processing_settings'],
                doc['created_at']
            ))
        
        return docs
    except Exception as e:
        st.error(f"Error fetching documents: {str(e)}")
        return []

def delete_document(user_id, doc_id):
    """Delete a document"""
    supabase = get_supabase_client()
    if not supabase:
        return False
    
    try:
        result = supabase.table('documents').delete().eq(
            'id', doc_id
        ).eq('user_id', user_id).execute()
        
        return len(result.data) > 0
    except Exception as e:
        st.error(f"Error deleting document: {str(e)}")
        return False

# Initialize database
if 'db_initialized' not in st.session_state:
    st.session_state.db_initialized = init_db()

# Streamlit page config
st.set_page_config(
    page_title="Gemini Vision OCR - Multi-Image",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS Styling
st.markdown("""
<style>
    .main-header {font-size: 2.5rem; color: #4285F4; text-align: center; margin-bottom: 0.5rem;}
    .sub-header {text-align: center; color: #666; margin-bottom: 2rem;}
    .success-box {padding: 1rem; border-radius: 0.5rem; background-color: #d4edda; border: 1px solid #c3e6cb; color: #155724; margin: 1rem 0;}
    .error-box {padding: 1rem; border-radius: 0.5rem; background-color: #f8d7da; border: 1px solid #f5c6cb; color: #721c24; margin: 1rem 0;}
    .warning-box {padding: 1rem; border-radius: 0.5rem; background-color: #fff3cd; border: 1px solid #ffecb5; color: #856404; margin: 1rem 0;}
    .info-box {padding: 1rem; border-radius: 0.5rem; background-color: #d1ecf1; border: 1px solid #bee5eb; color: #0c5460; margin: 1rem 0;}
    .stButton > button {width: 100%; background: linear-gradient(90deg, #4285F4, #34A853); color: white; border: none; border-radius: 0.5rem; padding: 0.5rem 1rem; font-weight: bold;}
    .image-container {border: 2px dashed #ccc; border-radius: 10px; padding: 10px; margin: 10px 0; text-align: center;}
    .processed-image {border: 2px solid #4CAF50; border-radius: 10px; padding: 10px; margin: 10px 0; background-color: #f0f8f0;}
    .user-info {background: linear-gradient(90deg, #4285F4, #34A853); color: white; padding: 0.5rem 1rem; border-radius: 0.5rem; margin-bottom: 1rem;}
    .document-card {border: 1px solid #ddd; border-radius: 0.5rem; padding: 1rem; margin: 0.5rem 0; background-color: #f9f9f9;}
    .formatting-panel {border: 2px solid #4285F4; border-radius: 0.5rem; padding: 1rem; margin: 1rem 0; background-color: #f8f9ff;}
</style>
""", unsafe_allow_html=True)

# Session state initialization
if 'authenticated' not in st.session_state:
    st.session_state.authenticated = False
if 'user_id' not in st.session_state:
    st.session_state.user_id = None
if 'username' not in st.session_state:
    st.session_state.username = None
if 'all_extracted_texts' not in st.session_state:
    st.session_state.all_extracted_texts = []
if 'processed_images' not in st.session_state:
    st.session_state.processed_images = []
if 'api_key_valid' not in st.session_state:
    st.session_state.api_key_valid = False
if 'show_formatting_panel' not in st.session_state:
    st.session_state.show_formatting_panel = False

# Authentication functions
def show_login():
    st.markdown('<h2 class="main-header">🔐 Login</h2>', unsafe_allow_html=True)
    
    with st.form("login_form"):
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        submitted = st.form_submit_button("Login")
        
        if submitted:
            if username and password:
                user = authenticate_user(username, password)
                if user:
                    st.session_state.authenticated = True
                    st.session_state.user_id = user[0]
                    st.session_state.username = user[1]
                    st.success(f"Welcome back, {username}!")
                    st.rerun()
                else:
                    st.error("Invalid username or password")
            else:
                st.error("Please enter both username and password")

def show_signup():
    st.markdown('<h2 class="main-header">📝 Sign Up</h2>', unsafe_allow_html=True)
    
    with st.form("signup_form"):
        username = st.text_input("Username")
        email = st.text_input("Email")
        password = st.text_input("Password", type="password")
        confirm_password = st.text_input("Confirm Password", type="password")
        submitted = st.form_submit_button("Sign Up")
        
        if submitted:
            if not all([username, email, password, confirm_password]):
                st.error("Please fill in all fields")
            elif password != confirm_password:
                st.error("Passwords do not match")
            elif len(password) < 6:
                st.error("Password must be at least 6 characters long")
            else:
                user_id = create_user(username, email, password)
                if user_id:
                    st.success("Account created successfully! Please login.")
                    st.session_state.show_login = True
                else:
                    st.error("Username or email already exists")

def show_supabase_setup():
    """Show Supabase setup instructions if needed"""
    st.markdown("""
    ## ⚠️ Database Setup Required
    
    Please create the following tables in your Supabase dashboard:
    
    ```sql
    -- Users table
    CREATE TABLE users (
        id SERIAL PRIMARY KEY,
        username VARCHAR(100) UNIQUE NOT NULL,
        email VARCHAR(255) UNIQUE NOT NULL,
        password_hash VARCHAR(64) NOT NULL,
        created_at TIMESTAMP DEFAULT NOW()
    );
    
    -- Documents table
    CREATE TABLE documents (
        id SERIAL PRIMARY KEY,
        user_id INTEGER NOT NULL,
        doc_name TEXT NOT NULL,
        extracted_texts TEXT NOT NULL,
        image_names TEXT NOT NULL,
        processing_settings TEXT NOT NULL,
        created_at TIMESTAMP DEFAULT NOW(),
        FOREIGN KEY (user_id) REFERENCES users (id) ON DELETE CASCADE
    );
    
    -- Create indexes
    CREATE INDEX idx_documents_user_id ON documents(user_id);
    CREATE INDEX idx_documents_created_at ON documents(created_at);
    ```
    
    Go to your Supabase project → SQL Editor → Run the above SQL commands.
    """)

# API key fetch
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')

# Check database initialization
if not st.session_state.db_initialized:
    st.error("Database connection failed or tables don't exist.")
    show_supabase_setup()
    st.stop()

# Authentication check
if not st.session_state.authenticated:
    st.markdown('<h1 class="main-header">🚀 Gemini Vision OCR - Multi-Image</h1>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">Please login or sign up to continue</p>', unsafe_allow_html=True)
    
    tab1, tab2 = st.tabs(["Login", "Sign Up"])
    
    with tab1:
        show_login()
    
    with tab2:
        show_signup()
        
else:
    # Main application for authenticated users
    st.markdown('<h1 class="main-header">🚀 Gemini Vision OCR - Multi-Image</h1>', unsafe_allow_html=True)
    
    # User info bar
    col1, col2 = st.columns([3, 1])
    with col1:
        st.markdown(f'<div class="user-info">Welcome, {st.session_state.username}!</div>', unsafe_allow_html=True)
    with col2:
        if st.button("Logout"):
            st.session_state.authenticated = False
            st.session_state.user_id = None
            st.session_state.username = None
            st.session_state.all_extracted_texts = []
            st.session_state.processed_images = []
            st.rerun()
    
    # Main tabs
    tab1, tab2, tab3 = st.tabs(["📷 New OCR", "📄 My Documents", "⚙️ Settings"])
    
    with tab3:  # Settings tab
        st.header("Settings")
        
        # API Key setup
        if not GEMINI_API_KEY:
            st.markdown("""
            <div class="warning-box">
                <strong>Setup Required:</strong> Please provide your Gemini API key.
            </div>
            """, unsafe_allow_html=True)
            temp_api_key = st.text_input("Enter Gemini API key:", type="password")
            if temp_api_key:
                GEMINI_API_KEY = temp_api_key
                st.success("API key loaded for this session!")
        else:
            st.markdown('<div class="success-box">API key loaded from environment</div>', unsafe_allow_html=True)
        
        if GEMINI_API_KEY:
            if st.button("Test API Key"):
                try:
                    genai.configure(api_key=GEMINI_API_KEY)
                    try:
                        model = genai.GenerativeModel("gemini-2.0-flash-exp")
                        response = model.generate_content("Test message. Respond with 'API works!'")
                    except Exception as model_error:
                        try:
                            model = genai.GenerativeModel("gemini-2.5-flash")
                            response = model.generate_content("Test message. Respond with 'API works!'")
                        except Exception:
                            model = genai.GenerativeModel("gemini-2.5-flash-lite")
                            response = model.generate_content("Test message. Respond with 'API works!'")
                    
                    if response and response.text and "API works" in response.text:
                        st.session_state.api_key_valid = True
                        st.success("API key is working!")
                    else:
                        st.warning("API responded but may have issues")
                except Exception as e:
                    st.session_state.api_key_valid = False
                    st.error(f"API test failed: {str(e)}")
    
    with tab2:  # My Documents tab
        st.header("My Documents")
        
        docs = get_user_documents(st.session_state.user_id)
        
        if not docs:
            st.markdown('<div class="info-box">No documents found. Process some images in the OCR tab to get started!</div>', unsafe_allow_html=True)
        else:
            for doc in docs:
                doc_id, doc_name, extracted_texts_json, image_names_json, settings_json, created_at = doc
                
                with st.expander(f"📄 {doc_name} - {created_at}"):
                    extracted_texts = json.loads(extracted_texts_json)
                    image_names = json.loads(image_names_json)
                    settings = json.loads(settings_json)
                    
                    st.write(f"**Images processed:** {len(image_names)}")
                    st.write(f"**Language:** {settings.get('target_language', 'N/A')}")
                    st.write(f"**OCR Mode:** {settings.get('ocr_mode', 'N/A')}")
                    
                    # Show extracted texts
                    for idx, text in enumerate(extracted_texts):
                        st.subheader(f"Image {idx + 1}: {image_names[idx]}")
                        st.text_area(f"Text {idx + 1}", value=text, height=100, key=f"saved_text_{doc_id}_{idx}")
                    
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        # Download as DOCX with formatting
                        if st.button(f"Download DOCX", key=f"docx_{doc_id}"):
                            try:
                                formatted_doc = create_formatted_document(
                                    doc_name, extracted_texts, image_names, settings,
                                    font_name="Calibri", font_size=11, line_spacing=1.15,
                                    margin_top=1.0, margin_bottom=1.0, margin_left=1.0, margin_right=1.0,
                                    separate_pages=settings.get('separate_pages', True),
                                    include_images=settings.get('include_images_in_docx', True)
                                )
                                
                                doc_buffer = io.BytesIO()
                                formatted_doc.save(doc_buffer)
                                doc_buffer.seek(0)
                                
                                st.download_button(
                                    label="Download DOCX",
                                    data=doc_buffer.getvalue(),
                                    file_name=f"{doc_name}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"download_docx_{doc_id}"
                                )
                            except Exception as e:
                                st.error(f"Error creating DOCX: {str(e)}")
                    
                    with col2:
                        # Download as TXT
                        if st.button(f"Download TXT", key=f"txt_{doc_id}"):
                            txt_content = f"Document: {doc_name}\n"
                            txt_content += f"Created: {created_at}\n"
                            txt_content += f"Language: {settings.get('target_language', 'N/A')}\n"
                            txt_content += f"OCR Mode: {settings.get('ocr_mode', 'N/A')}\n"
                            txt_content += "-" * 60 + "\n\n"
                            
                            for idx, text in enumerate(extracted_texts):
                                txt_content += f"Image {idx + 1}: {image_names[idx]}\n"
                                txt_content += text + "\n\n"
                            
                            st.download_button(
                                label="Download TXT",
                                data=txt_content,
                                file_name=f"{doc_name}.txt",
                                mime="text/plain",
                                key=f"download_txt_{doc_id}"
                            )
                    
                    with col3:
                        # Delete document
                        if st.button(f"🗑️ Delete", key=f"delete_{doc_id}", type="secondary"):
                            delete_document(st.session_state.user_id, doc_id)
                            st.success("Document deleted!")
                            st.rerun()

    def create_formatted_document(doc_name, extracted_texts, image_names, settings, 
                                font_name="Calibri", font_size=11, line_spacing=1.15,
                                margin_top=1.0, margin_bottom=1.0, margin_left=1.0, margin_right=1.0,
                                separate_pages=True, include_images=True, preserve_original_formatting=True,
                                add_page_numbers=False, **kwargs):
        """Create a formatted Word document"""
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(margin_top)
            section.bottom_margin = Inches(margin_bottom)
            section.left_margin = Inches(margin_left)
            section.right_margin = Inches(margin_right)
            
            # Add page numbers if requested
            if add_page_numbers:
                try:
                    from docx.oxml.shared import qn
                    from docx.oxml import parse_xml
                    footer = section.footer
                    footer_para = footer.paragraphs[0]
                    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add page number field
                    run = footer_para.runs[0] if footer_para.runs else footer_para.add_run()
                    fldChar1 = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                    instrText = parse_xml(r'<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE </w:instrText>')
                    fldChar2 = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                    
                    run._r.append(fldChar1)
                    run._r.append(instrText)
                    run._r.append(fldChar2)
                    
                    # Style the page number
                    run.font.name = font_name
                    run.font.size = Pt(font_size - 2)
                except Exception as e:
                    # If page numbering fails, continue without it
                    pass
        
        # Create custom styles
        styles = doc.styles
        
        # Title style
        if 'CustomTitle' not in [style.name for style in styles]:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = font_name
            title_font.size = Pt(font_size + 4)
            title_font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        
        # Heading style
        if 'CustomHeading' not in [style.name for style in styles]:
            heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
            heading_font = heading_style.font
            heading_font.name = font_name
            heading_font.size = Pt(font_size + 2)
            heading_font.bold = True
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
        
        # Body style
        if 'CustomBody' not in [style.name for style in styles]:
            body_style = styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = font_name
            body_font.size = Pt(font_size)
            body_style.paragraph_format.line_spacing = line_spacing
            body_style.paragraph_format.space_after = Pt(6)
        
        # Add title
        title = doc.add_paragraph('Gemini Vision OCR - Multi-Image', style='CustomTitle')
        
        # Add metadata
        doc.add_paragraph(f"Document: {doc_name}", style='CustomBody')
        doc.add_paragraph(f"Extraction Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", style='CustomBody')
        doc.add_paragraph(f"Language: {settings.get('target_language', 'N/A')}", style='CustomBody')
        doc.add_paragraph(f"OCR Mode: {settings.get('ocr_mode', 'N/A')}", style='CustomBody')
        doc.add_paragraph("", style='CustomBody')  # Empty line
        
        # Add content for each image
        for idx, text in enumerate(extracted_texts):
            if separate_pages and idx > 0:
                doc.add_page_break()
            
            # Add image heading
            doc.add_heading(f"Image {idx + 1}: {image_names[idx]}", level=1).style = 'CustomHeading'
            
            # Add image if requested and available
            if include_images and idx < len(st.session_state.processed_images):
                try:
                    img_stream = io.BytesIO()
                    st.session_state.processed_images[idx][1].save(img_stream, format='PNG')
                    img_stream.seek(0)
                    
                    paragraph = doc.add_paragraph()
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.add_picture(img_stream, width=Inches(4))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph("", style='CustomBody')  # Empty line after image
                except Exception as e:
                    doc.add_paragraph(f"[Could not embed image: {str(e)}]", style='CustomBody')
            
            # Add extracted text
            if preserve_original_formatting and text:
                # Split text into paragraphs and preserve formatting
                paragraphs = text.split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        para = doc.add_paragraph(style='CustomBody')
                        
                        # Check for special formatting
                        if para_text.strip().isupper() and len(para_text.strip()) < 100:
                            # Likely a heading - make it bold
                            run = para.add_run(para_text)
                            run.bold = True
                        elif para_text.startswith('    ') or para_text.startswith('\t'):
                            # Indented text - preserve indentation
                            run = para.add_run(para_text)
                        else:
                            run = para.add_run(para_text)
                    else:
                        doc.add_paragraph("", style='CustomBody')  # Empty line
            else:
                # Simple paragraph addition
                doc.add_paragraph(text, style='CustomBody')
            
            # Add spacing between sections
            if idx < len(extracted_texts) - 1:
                doc.add_paragraph("", style='CustomBody')
        
        return doc
    
    def show_formatting_panel():
        """Show the formatting panel for DOCX customization"""
        st.markdown('<div class="formatting-panel">', unsafe_allow_html=True)
        st.subheader("📝 Document Formatting Options")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.write("**Font Settings**")
            font_name = st.selectbox("Font Family", 
                ["Calibri", "Arial", "Times New Roman", "Georgia", "Verdana", "Cambria", "Tahoma"], 
                index=0, key="format_font")
            font_size = st.slider("Font Size", 8, 24, 11, key="format_size")
            line_spacing = st.selectbox("Line Spacing", 
                [1.0, 1.15, 1.5, 2.0], 
                index=1, key="format_spacing")
        
        with col2:
            st.write("**Page Settings**")
            margin_top = st.slider("Top Margin (inches)", 0.5, 2.0, 1.0, 0.1, key="margin_top")
            margin_bottom = st.slider("Bottom Margin (inches)", 0.5, 2.0, 1.0, 0.1, key="margin_bottom")
            margin_left = st.slider("Left Margin (inches)", 0.5, 2.0, 1.0, 0.1, key="margin_left")
            margin_right = st.slider("Right Margin (inches)", 0.5, 2.0, 1.0, 0.1, key="margin_right")
        
        with col3:
            st.write("**Document Options**")
            separate_pages = st.checkbox("Each image on new page", value=True, key="format_separate")
            include_images = st.checkbox("Include images in document", value=True, key="format_images")
            preserve_original_formatting = st.checkbox("Preserve original formatting", value=True, key="format_preserve")
            add_page_numbers = st.checkbox("Add page numbers", value=False, key="format_page_nums")
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        return {
            'font_name': font_name,
            'font_size': font_size,
            'line_spacing': line_spacing,
            'margin_top': margin_top,
            'margin_bottom': margin_bottom,
            'margin_left': margin_left,
            'margin_right': margin_right,
            'separate_pages': separate_pages,
            'include_images': include_images,
            'preserve_original_formatting': preserve_original_formatting,
            'add_page_numbers': add_page_numbers
        }

    with tab1:  # OCR Processing tab
        if not GEMINI_API_KEY:
            st.markdown("""
            <div class="warning-box">
                <strong>Setup Required:</strong> Please configure your API key in the Settings tab first.
            </div>
            """, unsafe_allow_html=True)
        else:
            # Sidebar options
            with st.sidebar:
                st.header("OCR Configuration")
                
                target_language = st.selectbox(
                    "Target Language:",
                    ["Same as original", "English", "Spanish", "French", "German", "Italian",
                     "Portuguese", "Russian", "Chinese", "Japanese", "Korean",
                     "Arabic", "Hindi", "Dutch", "Swedish", "Norwegian"],
                    index=1
                )
                ocr_mode = st.selectbox(
                    "OCR Mode:",
                    ["Legal/Official Document", "Handwriting Focus", "Mixed Text", "Document Scan", "Creative/Artistic Text"],
                    index=0
                )
                st.divider()

                st.subheader("Processing Options")
                preserve_formatting = st.checkbox("Preserve text formatting", value=True)
                include_images_in_docx = st.checkbox("Include images in DOCX", value=True)
                st.divider()

                st.subheader("Multi-Image Options")
                process_all_at_once = st.checkbox("Process all images at once", value=False)
                separate_pages = st.checkbox("Each image on new page", value=True)
                st.divider()

                if st.button("Clear Current Session"):
                    st.session_state.all_extracted_texts = []
                    st.session_state.processed_images = []
                    st.rerun()

            def clean_extracted_text(text):
                """Clean and format the extracted text properly"""
                if not text:
                    return text
                
                import re
                
                # Replace underline tags with underscore formatting
                text = re.sub(r'<u>(.*?)</u>', r'_\1_', text, flags=re.IGNORECASE)
                text = re.sub(r'<b>(.*?)</b>', r'**\1**', text, flags=re.IGNORECASE)
                text = re.sub(r'<i>(.*?)</i>', r'*\1*', text, flags=re.IGNORECASE)
                
                # Remove any remaining HTML tags
                text = re.sub(r'<[^>]+>', '', text)
                
                # Fix multiple spaces but preserve intentional indentation
                lines = text.split('\n')
                cleaned_lines = []
                
                for line in lines:
                    leading_spaces = len(line) - len(line.lstrip())
                    content = line.strip()
                    if content:
                        cleaned_lines.append(' ' * leading_spaces + content)
                    else:
                        cleaned_lines.append('')
                
                return '\n'.join(cleaned_lines)

            def get_model():
                """Get the best available Gemini model"""
                genai.configure(api_key=GEMINI_API_KEY)
                
                models_to_try = [
                    "gemini-2.0-flash-exp",
                    "gemini-2.5-flash",
                    "gemini-2.5-flash-lite",
                ]
                
                for model_name in models_to_try:
                    try:
                        return genai.GenerativeModel(model_name)
                    except Exception:
                        continue
                
                raise Exception("No compatible Gemini model found. Please check your API access and model availability.")

            # Main App Area
            st.header("Upload Multiple Images")

            uploaded_files = st.file_uploader(
                "Upload images...",
                type=['png', 'jpg', 'jpeg', 'bmp', 'tiff', 'webp'],
                accept_multiple_files=True
            )

            if uploaded_files:
                st.success(f"{len(uploaded_files)} images uploaded!")
                st.subheader("Uploaded Images")
                cols = st.columns(min(3, len(uploaded_files)))
                for idx, uploaded_file in enumerate(uploaded_files):
                    with cols[idx % 3]:
                        image = Image.open(uploaded_file)
                        st.image(image, caption=uploaded_file.name, width=300)
                        st.caption(f"Size: {image.size[0]}x{image.size[1]}px")

                def create_prompt():
                    if ocr_mode == "Legal/Official Document":
                        base_prompt = ("You are an expert at extracting text from legal and official documents. Extract ALL text with perfect accuracy, "
                                       "maintaining exact formatting, numbering systems, indentation, and legal document structure. "
                                       "Pay special attention to section numbers, subsections, clauses, and hierarchical organization.")
                    elif ocr_mode == "Handwriting Focus":
                        base_prompt = ("You are an expert at reading handwritten text. Extract ALL text from this handwritten image with high accuracy. "
                                       "Pay special attention to cursive writing, connected letters, and personal writing styles.")
                    elif ocr_mode == "Mixed Text":
                        base_prompt = "Extract all text from this image, which may contain both handwritten and printed text."
                    elif ocr_mode == "Document Scan":
                        base_prompt = "Extract all text from this document image, maintaining the exact structure and layout."
                    else:
                        base_prompt = "Extract text from this image, which may contain artistic, stylized, or decorative text."

                    if target_language != "Same as original":
                        base_prompt += f" Translate the extracted text to {target_language}."
                    
                    if preserve_formatting:
                        base_prompt += """ CRITICAL FORMATTING RULES:
- Preserve ALL numbering exactly: (6), (7), (8), etc.
- Use proper indentation with spaces (not tabs)
- For underlined text: Use CAPITAL LETTERS or add underscores like _HEADING_
- DO NOT use HTML tags like <u>, <b>, or similar
- Maintain exact spacing and line breaks as shown
- Keep the same paragraph indentation levels
- Preserve all punctuation and special characters exactly
- Use plain text formatting only - no markup tags
- For emphasized text, use CAPITALS or _underscores_
- Maintain the visual hierarchy with proper spacing between sections"""
                    else:
                        base_prompt += " Preserve basic paragraph structure and line breaks."
                        
                    base_prompt += "\n\nProvide ONLY the extracted text in plain text format. NO HTML tags, NO markup. Use spaces for indentation and underscores or CAPITALS for emphasis."
                    return base_prompt

                st.header("Process Images")
                col1, col2 = st.columns(2)

                # Processing Column
                with col1:
                    if process_all_at_once:
                        if st.button("Process All Images"):
                            try:
                                with st.spinner("Analyzing images..."):
                                    model = get_model()
                                    prompt = create_prompt()
                                    st.session_state.all_extracted_texts = []
                                    st.session_state.processed_images = []
                                    for uploaded_file in uploaded_files:
                                        image = Image.open(uploaded_file)
                                        response = model.generate_content([prompt, image])
                                        text = response.text.strip() if response.text else ""
                                        text = clean_extracted_text(text)
                                        if text:
                                            st.session_state.all_extracted_texts.append(text)
                                            st.session_state.processed_images.append((uploaded_file.name, image))
                                            st.markdown(f'<div class="success-box">Text extracted from {uploaded_file.name}</div>', unsafe_allow_html=True)
                                        else:
                                            st.markdown(f'<div class="error-box">No text detected in {uploaded_file.name}</div>', unsafe_allow_html=True)
                            except Exception as e:
                                st.error(f"Error: {str(e)}")
                    else:
                        selected_image = st.selectbox("Select an image to process:", [f.name for f in uploaded_files])
                        if st.button(f"Process {selected_image}"):
                            try:
                                with st.spinner(f"Analyzing {selected_image}..."):
                                    model = get_model()
                                    prompt = create_prompt()
                                    selected_file = next(f for f in uploaded_files if f.name == selected_image)
                                    image = Image.open(selected_file)
                                    response = model.generate_content([prompt, image])
                                    text = response.text.strip() if response.text else ""
                                    text = clean_extracted_text(text)
                                    if text:
                                        st.session_state.all_extracted_texts.append(text)
                                        st.session_state.processed_images.append((selected_file.name, image))
                                        st.markdown(f'<div class="success-box">Text extracted from {selected_file.name}</div>', unsafe_allow_html=True)
                                    else:
                                        st.markdown(f'<div class="error-box">No text detected in {selected_file.name}</div>', unsafe_allow_html=True)
                            except Exception as e:
                                st.error(f"Error: {str(e)}")

                # Results Column
                with col2:
                    if st.session_state.all_extracted_texts:
                        st.header("Extracted Text")
                        for idx, text in enumerate(st.session_state.all_extracted_texts):
                            st.subheader(f"Image {idx + 1}: {st.session_state.processed_images[idx][0]}")
                            st.text_area(f"Text from {st.session_state.processed_images[idx][0]}", value=text, height=150, key=f"text_{idx}")

                        st.header("Save & Download")
                        
                        # Document name input
                        doc_name = st.text_input("Document Name:", value=f"OCR_Document_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
                        
                        if st.button("Save Document"):
                            if doc_name.strip():
                                try:
                                    # Prepare data for saving
                                    image_names = [img[0] for img in st.session_state.processed_images]
                                    settings = {
                                        'target_language': target_language,
                                        'ocr_mode': ocr_mode,
                                        'preserve_formatting': preserve_formatting,
                                        'include_images_in_docx': include_images_in_docx,
                                        'separate_pages': separate_pages
                                    }
                                    
                                    doc_id = save_document(
                                        st.session_state.user_id,
                                        doc_name.strip(),
                                        st.session_state.all_extracted_texts,
                                        image_names,
                                        settings
                                    )
                                    
                                    if doc_id:
                                        st.success(f"Document '{doc_name}' saved successfully!")
                                        st.markdown('<div class="info-box">You can view and download your saved documents from the "My Documents" tab.</div>', unsafe_allow_html=True)
                                    else:
                                        st.error("Failed to save document. Please try again.")
                                    
                                except Exception as e:
                                    st.error(f"Error saving document: {str(e)}")
                            else:
                                st.error("Please enter a document name")

                        st.divider()
                        
                        # Formatting Panel Toggle
                        if st.button("📝 Advanced Formatting Options"):
                            st.session_state.show_formatting_panel = not st.session_state.show_formatting_panel
                        
                        # Show formatting panel if enabled
                        format_settings = None
                        if st.session_state.show_formatting_panel:
                            format_settings = show_formatting_panel()
                        
                        # Quick download options
                        if st.button("Quick Download as DOCX"):
                            try:
                                if format_settings:
                                    # Use custom formatting
                                    formatted_doc = create_formatted_document(
                                        doc_name or f"OCR_Document_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                                        st.session_state.all_extracted_texts,
                                        [img[0] for img in st.session_state.processed_images],
                                        {'target_language': target_language, 'ocr_mode': ocr_mode},
                                        **format_settings
                                    )
                                else:
                                    # Use default formatting
                                    formatted_doc = create_formatted_document(
                                        doc_name or f"OCR_Document_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                                        st.session_state.all_extracted_texts,
                                        [img[0] for img in st.session_state.processed_images],
                                        {'target_language': target_language, 'ocr_mode': ocr_mode},
                                        font_name="Calibri", font_size=11, line_spacing=1.15,
                                        margin_top=1.0, margin_bottom=1.0, margin_left=1.0, margin_right=1.0,
                                        separate_pages=separate_pages, include_images=include_images_in_docx
                                    )
                                
                                doc_buffer = io.BytesIO()
                                formatted_doc.save(doc_buffer)
                                doc_buffer.seek(0)
                                
                                st.download_button(
                                    label="Download Formatted DOCX",
                                    data=doc_buffer.getvalue(),
                                    file_name=f"formatted_{doc_name or 'extracted_text'}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                            except Exception as e:
                                st.error(f"Error creating DOCX: {str(e)}")

                        if st.button("Quick Download as TXT"):
                            txt_content = "Gemini Vision OCR - Multi-Image\n"
                            txt_content += f"Extraction Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
                            txt_content += f"Target Language: {target_language}\n"
                            txt_content += f"OCR Mode: {ocr_mode}\n"
                            txt_content += "-" * 60 + "\n\n"
                            for idx, (filename, _) in enumerate(st.session_state.processed_images):
                                txt_content += f"Image {idx + 1}: {filename}\n"
                                txt_content += st.session_state.all_extracted_texts[idx] + "\n\n"
                                
                            st.download_button(
                                label="Download TXT",
                                data=txt_content,
                                file_name=f"extracted_text_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                                mime="text/plain"
                            )